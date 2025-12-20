"""
Document text extraction service using proven libraries.

This module provides robust text extraction from various document formats
using industry-standard libraries:
- pdfplumber: Primary PDF extraction (most reliable for structured documents)
- PyPDF2: Fallback PDF extraction
- python-docx: Microsoft Word document extraction
"""
import logging
from typing import Dict, Any, Optional
from pathlib import Path
import io

# PDF extraction
import pdfplumber
import PyPDF2

# DOCX extraction
from docx import Document

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Professional-grade document text extraction service.
    
    Supports: PDF, DOCX, TXT
    Uses proven, officially maintained libraries for maximum reliability.
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    def __init__(self):
        """Initialize document extractor"""
        self.extraction_stats = {
            'total_extracted': 0,
            'successful': 0,
            'failed': 0
        }
    
    async def extract_text(
        self,
        file_content: bytes,
        filename: str,
        use_fallback: bool = True
    ) -> Dict[str, Any]:
        """
        Extract text from document with metadata.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename with extension
            use_fallback: Use fallback extraction methods if primary fails
            
        Returns:
            Dictionary containing:
                - text: Extracted text content
                - metadata: Document metadata (pages, encoding, etc.)
                - extraction_method: Method used for extraction
                - success: Boolean indicating success
                - error: Error message if failed
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            self.extraction_stats['failed'] += 1
            self.extraction_stats['total_extracted'] += 1
            return {
                'text': '',
                'metadata': {},
                'extraction_method': 'none',
                'success': False,
                'error': f'Unsupported file type: {file_ext}'
            }
        
        logger.info(f"Extracting text from {filename} ({file_ext})")
        
        try:
            if file_ext == '.pdf':
                result = await self._extract_from_pdf(file_content, use_fallback)
            elif file_ext in ['.docx', '.doc']:
                result = await self._extract_from_docx(file_content)
            elif file_ext == '.txt':
                result = await self._extract_from_txt(file_content)
            else:
                result = {
                    'text': '',
                    'metadata': {},
                    'extraction_method': 'none',
                    'success': False,
                    'error': f'Unhandled file type: {file_ext}'
                }
            
            if result['success']:
                self.extraction_stats['successful'] += 1
            else:
                self.extraction_stats['failed'] += 1
            
            self.extraction_stats['total_extracted'] += 1
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            self.extraction_stats['failed'] += 1
            self.extraction_stats['total_extracted'] += 1
            
            return {
                'text': '',
                'metadata': {},
                'extraction_method': 'none',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_from_pdf(
        self,
        file_content: bytes,
        use_fallback: bool = True
    ) -> Dict[str, Any]:
        """
        Extract text from PDF using pdfplumber (primary) or PyPDF2 (fallback).
        
        pdfplumber is more reliable for structured documents and maintains
        better layout information.
        """
        # Try pdfplumber first (most reliable)
        try:
            text_parts = []
            metadata = {
                'pages': 0,
                'extraction_method': 'pdfplumber'
            }
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Also extract metadata from first page
                    if page_num == 1 and pdf.metadata:
                        metadata.update({
                            'title': pdf.metadata.get('Title', ''),
                            'author': pdf.metadata.get('Author', ''),
                            'subject': pdf.metadata.get('Subject', ''),
                            'creator': pdf.metadata.get('Creator', ''),
                        })
            
            full_text = '\n\n'.join(text_parts)
            
            if full_text.strip():
                logger.info(f"Successfully extracted {len(full_text)} characters using pdfplumber")
                return {
                    'text': full_text,
                    'metadata': metadata,
                    'extraction_method': 'pdfplumber',
                    'success': True,
                    'error': None
                }
            else:
                raise ValueError("No text extracted from PDF")
                
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
            
            # Fallback to PyPDF2
            if use_fallback:
                try:
                    return await self._extract_from_pdf_pypdf2(file_content)
                except Exception as fallback_error:
                    logger.error(f"PyPDF2 fallback also failed: {str(fallback_error)}")
                    return {
                        'text': '',
                        'metadata': {},
                        'extraction_method': 'none',
                        'success': False,
                        'error': f"Both extraction methods failed. Primary: {str(e)}, Fallback: {str(fallback_error)}"
                    }
            else:
                return {
                    'text': '',
                    'metadata': {},
                    'extraction_method': 'none',
                    'success': False,
                    'error': str(e)
                }
    
    async def _extract_from_pdf_pypdf2(self, file_content: bytes) -> Dict[str, Any]:
        """Fallback PDF extraction using PyPDF2"""
        logger.info("Attempting PyPDF2 extraction as fallback")
        
        text_parts = []
        metadata = {
            'extraction_method': 'pypdf2'
        }
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        metadata['pages'] = len(pdf_reader.pages)
        
        # Extract metadata
        if pdf_reader.metadata:
            metadata.update({
                'title': pdf_reader.metadata.get('/Title', ''),
                'author': pdf_reader.metadata.get('/Author', ''),
                'subject': pdf_reader.metadata.get('/Subject', ''),
                'creator': pdf_reader.metadata.get('/Creator', ''),
            })
        
        # Extract text from all pages
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        full_text = '\n\n'.join(text_parts)
        
        if not full_text.strip():
            raise ValueError("No text extracted using PyPDF2")
        
        logger.info(f"Successfully extracted {len(full_text)} characters using PyPDF2")
        
        return {
            'text': full_text,
            'metadata': metadata,
            'extraction_method': 'pypdf2',
            'success': True,
            'error': None
        }
    
    async def _extract_from_docx(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from DOCX using python-docx.
        
        python-docx is the official library for reading Microsoft Word documents.
        """
        try:
            document = Document(io.BytesIO(file_content))
            
            # Extract all paragraphs
            text_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = '\n\n'.join(text_parts)
            
            metadata = {
                'paragraphs': len(document.paragraphs),
                'tables': len(document.tables),
                'extraction_method': 'python-docx'
            }
            
            # Extract core properties if available
            if hasattr(document, 'core_properties'):
                props = document.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': str(props.created) if props.created else '',
                    'modified': str(props.modified) if props.modified else '',
                })
            
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            
            return {
                'text': full_text,
                'metadata': metadata,
                'extraction_method': 'python-docx',
                'success': True,
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return {
                'text': '',
                'metadata': {},
                'extraction_method': 'none',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_from_txt(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from TXT files with encoding detection.
        
        Handles various encodings commonly found in text files.
        """
        # Try common encodings
        encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings_to_try:
            try:
                text = file_content.decode(encoding)
                
                logger.info(f"Successfully decoded TXT file using {encoding}")
                
                return {
                    'text': text,
                    'metadata': {
                        'encoding': encoding,
                        'size_bytes': len(file_content),
                        'extraction_method': 'decode'
                    },
                    'extraction_method': 'decode',
                    'success': True,
                    'error': None
                }
            except (UnicodeDecodeError, AttributeError):
                continue
        
        # If all encodings fail
        logger.error("Failed to decode TXT file with any known encoding")
        return {
            'text': '',
            'metadata': {},
            'extraction_method': 'none',
            'success': False,
            'error': 'Unable to decode file with any known encoding'
        }
    
    def get_stats(self) -> Dict[str, int]:
        """Get extraction statistics"""
        return self.extraction_stats.copy()
